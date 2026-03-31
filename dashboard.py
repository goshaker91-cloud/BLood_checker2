import streamlit as st
import pandas as pd
import sqlite3
import datetime
import requests
import base64
import re
from docx import Document

# ========== CONFIGURATION ==========
# Get API key from Streamlit secrets
CLINIKO_API_KEY = st.secrets["CLINIKO_API_KEY"]
CLINIKO_SHARD = "au1"                   # Change if your shard is eu1, us1, etc.
CLINIKO_BASE_URL = f"https://api.{CLINIKO_SHARD}.cliniko.com/v1"

# ========== HELPER FUNCTIONS ==========

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_patient_info(text):
    """Extract name, DOB, NHS number from document text."""
    name = None
    dob = None
    nhs = None
    name_match = re.search(r"Patient\s*[:\-]?\s*([A-Za-z\s]+)", text, re.IGNORECASE)
    if name_match:
        name = name_match.group(1).strip()
    dob_match = re.search(r"DOB\s*[:\-]?\s*(\d{2}[/\-]\d{2}[/\-]\d{4})", text, re.IGNORECASE)
    if dob_match:
        dob = dob_match.group(1)
    nhs_match = re.search(r"NHS\s*[:\-]?\s*(\d{10})", text, re.IGNORECASE)
    if nhs_match:
        nhs = nhs_match.group(1)
    return {"name": name, "dob": dob, "nhs": nhs}

def extract_results(docx_path):
    doc = Document(docx_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    pattern = r"^([AB])\s+([\w\s\-/]+?)\s+([\d\.]+)\s+([\w/]+)\s+(.+)$"
    results = []
    for line in text.split("\n"):
        line = line.strip()
        match = re.match(pattern, line)
        if match:
            test_name = match.group(2).strip()
            result = match.group(3)
            unit = match.group(4)
            ref_range = match.group(5)
            abnormal = "**" in line or "Low" in ref_range or "High" in ref_range
            results.append({
                "test": test_name,
                "result": result,
                "unit": unit,
                "reference": ref_range,
                "abnormal": abnormal
            })
    return results

def search_cliniko_patient(nhs_number=None, first_name=None, last_name=None):
    credentials = f"{CLINIKO_API_KEY}:"
    encoded = base64.b64encode(credentials.encode()).decode()
    headers = {
        "Authorization": f"Basic {encoded}",
        "Accept": "application/json",
        "User-Agent": "Blood Results Checker (your-email@example.com)"
    }
    params = []
    if nhs_number:
        params.append(f"q[]=nhs_number:{nhs_number}")
    elif first_name and last_name:
        params.append(f"q[]=first_name:{first_name}")
        params.append(f"q[]=last_name:{last_name}")
    else:
        return None
    url = f"{CLINIKO_BASE_URL}/patients"
    if params:
        url += "?" + "&".join(params)
    try:
        response = requests.get(url, headers=headers)
        if response.status_code == 200:
            data = response.json()
            patients = data.get("patients", [])
            if patients:
                return patients[0]
    except Exception as e:
        st.error(f"Cliniko error: {e}")
    return None

def create_cliniko_note(patient_id, note_content, note_type="Clinical note"):
    """
    Create a clinical note for a patient in Cliniko.
    """
    credentials = f"{CLINIKO_API_KEY}:"
    encoded = base64.b64encode(credentials.encode()).decode()
    headers = {
        "Authorization": f"Basic {encoded}",
        "Accept": "application/json",
        "Content-Type": "application/json",
        "User-Agent": "Blood Results Checker (your-email@example.com)"
    }
    payload = {
        "note": note_content,
        "type": note_type
    }
    url = f"{CLINIKO_BASE_URL}/patients/{patient_id}/notes"
    try:
        response = requests.post(url, headers=headers, json=payload)
        if response.status_code in [200, 201]:
            return True
        else:
            st.error(f"Failed to create note: {response.status_code} - {response.text}")
            return False
    except Exception as e:
        st.error(f"Error creating note: {e}")
        return False

def init_db():
    conn = sqlite3.connect('blood_reviews.db')
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS reviews (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            patient_name TEXT,
            patient_nhs TEXT,
            upload_time TEXT,
            abnormal_count INTEGER,
            plan TEXT,
            confirmed INTEGER DEFAULT 0
        )
    ''')
    conn.commit()
    conn.close()

def save_review(patient, abnormal_count, plan):
    conn = sqlite3.connect('blood_reviews.db')
    c = conn.cursor()
    c.execute('''
        INSERT INTO reviews (patient_name, patient_nhs, upload_time, abnormal_count, plan)
        VALUES (?, ?, ?, ?, ?)
    ''', (patient.get('name'), patient.get('nhs'), datetime.datetime.now().isoformat(), abnormal_count, plan))
    conn.commit()
    conn.close()

# ========== INITIALISE DATABASE ==========
init_db()

# ========== STREAMLIT UI ==========
st.set_page_config(page_title="Blood Results Checker", layout="wide")
st.title("🩸 Blood Results Checker")

uploaded_file = st.file_uploader("Upload a blood test document", type=["docx"])

if uploaded_file:
    with open("temp.docx", "wb") as f:
        f.write(uploaded_file.getbuffer())

    # Extract patient info and results
    text = extract_text_from_docx("temp.docx")
    patient_info = extract_patient_info(text)
    results = extract_results("temp.docx")
    abnormal = [r for r in results if r["abnormal"]]

    # Try to find patient in Cliniko
    if patient_info['nhs']:
        cliniko_patient = search_cliniko_patient(nhs_number=patient_info['nhs'])
    elif patient_info['name']:
        parts = patient_info['name'].split(maxsplit=1)
        first = parts[0] if parts else None
        last = parts[1] if len(parts) > 1 else None
        cliniko_patient = search_cliniko_patient(first_name=first, last_name=last)
    else:
        cliniko_patient = None

    # Display patient match status
    if cliniko_patient:
        st.success(f"✅ Matched Cliniko patient: {cliniko_patient.get('first_name')} {cliniko_patient.get('last_name')} (ID: {cliniko_patient['id']})")
    else:
        st.warning("⚠️ No matching patient found in Cliniko")

    # Show all results
    st.subheader("All Results")
    df = pd.DataFrame(results)
    st.dataframe(df, use_container_width=True)

    # Show only abnormal results
    if abnormal:
        st.subheader("⚠️ Abnormal Results")
        st.dataframe(pd.DataFrame(abnormal), use_container_width=True)

    # Plan input and save
    plan = st.text_area("Management plan")
    if st.button("Save review"):
        save_review(patient_info, len(abnormal), plan)
        st.success("Review saved locally.")

    # New: Push to Cliniko button
    if st.button("Push plan to Cliniko"):
        if cliniko_patient:
            note_text = "Blood results review:\n"
            note_text += "Abnormal results:\n"
            for r in abnormal:
                note_text += f"- {r['test']}: {r['result']} {r['unit']} (ref {r['reference']})\n"
            note_text += f"\nPlan: {plan}"
            success = create_cliniko_note(cliniko_patient['id'], note_text)
            if success:
                st.success("Plan pushed to Cliniko as a clinical note.")
            else:
                st.error("Failed to push plan to Cliniko.")
        else:
            st.warning("No patient matched in Cliniko. Cannot push note.")